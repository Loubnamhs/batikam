# -*- coding: utf-8 -*-
"""Export DOCX — Batikam Rénove.

Design sobre et professionnel : palette minimaliste, tableau aéré,
typographie claire, totaux élégants.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from app.models.devis import Devis, Ligne
from app.services.branding import resolve_logo_path
from app.services.company_info import get_company_info
from app.services.paths import resolve_resource_path
from app.services.document_theme import (
    COLOR_BORDER,
    COLOR_GOLD,
    COLOR_NAVY,
    COLOR_SECTION,
    COLOR_TABLE_ALT_BG,
    COLOR_TABLE_GRID,
    COLOR_TABLE_HEADER_BG,
    COLOR_TABLE_HEADER_TEXT,
    COLOR_TABLE_LOT_BG,
    COLOR_TABLE_SUBTOTAL_BG,
    COLOR_TABLE_SUBTOTAL_TEXT,
    COLOR_TEXT,
    COLOR_MUTED,
    DOCX_FOOTER_DISTANCE_CM,
    DOCX_FOOTER_LINE1_AFTER_PT,
    DOCX_FOOTER_LINE1_FONT_PT,
    DOCX_FOOTER_LINE2_FONT_PT,
    DOCX_FOOTER_LINE_COLOR,
    DOCX_FOOTER_TEXT_COLOR,
    DOCX_HEADER_LEFT_COL_CM,
    DOCX_HEADER_RIGHT_COL_CM,
    DOCX_HEADER_LOGO_WIDTH_CM,
    DOCX_HEADER_DOC_BLOCK_GAP_PT,
    DOCX_HEADER_GAP_AFTER_CLIENT_PT,
    DOCX_HEADER_GAP_AFTER_AFFAIRE_PT,
    DOCX_LINE_TABLE_WIDTHS_CM,
    DOCX_CONTENT_TABLE_WIDTH_CM,
    DOCX_SPACE_BEFORE_TABLE_PT,
    LINE_TABLE_HEADERS,
)


def _set_cell_multiline(cell, text: str) -> None:
    lines = text.splitlines()
    if not lines:
        cell.text = ""
        return
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        cell.add_paragraph(line)


def euro_fr(value) -> str:
    return f"{float(value):,.2f} €".replace(",", "\u202f").replace(".", ",")


def date_fr(d) -> str:
    return d.strftime("%d/%m/%Y")


class DOCXExporter:

    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path or str(
            resolve_resource_path("templates", "devis_template.docx"))

    def _has_grouped_lots(self, devis: Devis) -> bool:
        if not devis.utiliser_lots:
            return False
        non_empty = [lot for lot in devis.lots if lot.lignes]
        if len(non_empty) <= 1:
            return bool((non_empty[0].nom if non_empty else "").strip())
        return True

    def _iter_all_lines(self, devis: Devis) -> list[Ligne]:
        lines: list[Ligne] = []
        for lot in devis.lots:
            lines.extend(lot.lignes)
        return lines

    def export(self, devis: Devis, output_path: str):
        self._export_with_python_docx(devis, output_path)

    def _export_with_python_docx(self, devis: Devis, output_path: str) -> None:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        company   = get_company_info()
        logo_path = resolve_logo_path()
        is_facture = (devis.statut or "").lower() == "facture"

        # Couleurs
        C_NAVY    = RGBColor.from_string(COLOR_NAVY.replace("#", ""))
        C_MUTED   = RGBColor.from_string(COLOR_MUTED.replace("#", ""))
        C_TEXT    = RGBColor.from_string(COLOR_TEXT.replace("#", ""))
        C_TH_TEXT = RGBColor.from_string(COLOR_TABLE_HEADER_TEXT.replace("#", ""))
        C_SUB_TXT = RGBColor.from_string(COLOR_TABLE_SUBTOTAL_TEXT.replace("#", ""))
        C_WHITE   = RGBColor(255, 255, 255)
        FOOTER_TEXT  = RGBColor.from_string(DOCX_FOOTER_TEXT_COLOR.replace("#", ""))
        FOOTER_LINE  = DOCX_FOOTER_LINE_COLOR.replace("#", "")

        HEX_GRID    = COLOR_TABLE_GRID.replace("#", "")
        HEX_TH_BG   = COLOR_TABLE_HEADER_BG.replace("#", "")
        HEX_LOT_BG  = COLOR_TABLE_LOT_BG.replace("#", "")
        HEX_SUB_BG  = COLOR_TABLE_SUBTOTAL_BG.replace("#", "")
        HEX_ALT_BG  = COLOR_TABLE_ALT_BG.replace("#", "")
        HEX_SECTION = COLOR_SECTION.replace("#", "")
        HEX_BORDER  = COLOR_BORDER.replace("#", "")
        HEX_NAVY    = COLOR_NAVY.replace("#", "")

        # ── Helpers XML ──────────────────────────────────────────────────────

        def _fix_table(tbl, total_cm: float) -> None:
            """Force largeur exacte + indentation 0 — remplace les éléments existants."""
            twips = int(total_cm / 2.54 * 1440)
            tbl_el = tbl._tbl
            tbl_pr = tbl_el.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl_el.insert(0, tbl_pr)

            def _replace(tag: str, attrs: dict) -> None:
                # Supprimer tous les éléments existants du même tag
                for old in tbl_pr.findall(qn(tag)):
                    tbl_pr.remove(old)
                el = OxmlElement(tag)
                for k, v in attrs.items():
                    el.set(qn(k), v)
                tbl_pr.append(el)

            _replace("w:tblW",          {"w:w": str(twips), "w:type": "dxa"})
            _replace("w:tblInd",        {"w:w": "0",        "w:type": "dxa"})
            _replace("w:tblCellSpacing",{"w:w": "0",        "w:type": "dxa"})

        def _shading(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tc_pr.append(shd)

        def _fix_cell_w(cell, width_cm: float) -> None:
            """Force w:tcW sur la cellule — Word respecte cette valeur à l'export PDF."""
            twips = int(width_cm / 2.54 * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(old)
            tcw = OxmlElement("w:tcW")
            tcw.set(qn("w:w"), str(twips))
            tcw.set(qn("w:type"), "dxa")
            tc_pr.append(tcw)

        def _borders(cell, color: str = HEX_GRID, size: str = "4",
                     sides=("top", "left", "bottom", "right")) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_b = OxmlElement("w:tcBorders")
            for side in sides:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), size)
                el.set(qn("w:color"), color)
                tc_b.append(el)
            tc_pr.append(tc_b)

        def _no_borders(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_b = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "nil")
                tc_b.append(el)
            tc_pr.append(tc_b)

        def _margins(cell, l=80, r=80, t=40, b=40) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_m = OxmlElement("w:tcMar")
            for side, val in (("left", l), ("right", r), ("top", t), ("bottom", b)):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), str(max(0, int(val))))
                el.set(qn("w:type"), "dxa")
                tc_m.append(el)
            tc_pr.append(tc_m)

        def _no_split(row) -> None:
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:cantSplit"))

        def _font(run, size: float, bold: bool = False,
                  color: Optional[RGBColor] = None, italic: bool = False) -> None:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            if color is not None:
                run.font.color.rgb = color

        def _spacing(p, before: float = 0, after: float = 3, line: float = 1.1) -> None:
            fmt = p.paragraph_format
            fmt.space_before = Pt(max(0.0, float(before)))
            fmt.space_after  = Pt(max(0.0, float(after)))
            fmt.line_spacing = line

        def _top_border_para(p, color: str, size: str = "6") -> None:
            p_pr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), size)
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), color)
            bdr.append(top)
            p_pr.append(bdr)

        def _bottom_border_para(p, color: str, size: str = "4") -> None:
            p_pr = p._p.get_or_add_pPr()
            bdr  = p_pr.find(qn("w:pBdr"))
            if bdr is None:
                bdr = OxmlElement("w:pBdr")
                p_pr.append(bdr)
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), size)
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), color)
            bdr.append(bot)

        # ── Document ─────────────────────────────────────────────────────────
        doc = Document()
        section = doc.sections[0]
        section.top_margin    = Cm(1.1)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(1.2)
        section.right_margin  = Cm(1.2)
        section.footer_distance = Cm(DOCX_FOOTER_DISTANCE_CM)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        # ── Header : logo gauche | infos document droite ───────────────────
        is_facture_flag = is_facture
        label    = "FACTURE" if is_facture_flag else "DEVIS"
        numero   = devis.numero or "—"
        d_cree   = date_fr(devis.date_devis)

        top = doc.add_table(rows=1, cols=2)
        top.autofit = False
        top.columns[0].width = Cm(DOCX_HEADER_LEFT_COL_CM)
        top.columns[1].width = Cm(DOCX_HEADER_RIGHT_COL_CM)
        _fix_table(top, DOCX_HEADER_LEFT_COL_CM + DOCX_HEADER_RIGHT_COL_CM)
        lc = top.cell(0, 0)
        rc = top.cell(0, 1)
        lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        rc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _no_borders(lc); _no_borders(rc)
        _margins(lc, 0, 0, 0, 0); _margins(rc, 0, 0, 0, 0)

        # Logo
        p_logo = lc.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p_logo, after=0)
        if logo_path and Path(logo_path).exists():
            p_logo.add_run().add_picture(str(logo_path), width=Cm(DOCX_HEADER_LOGO_WIDTH_CM))

        # Type de document — grand, ardoise
        p = rc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, after=2)
        _font(p.add_run(label), 22, bold=True, color=C_NAVY)

        # N° discret
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, after=6)
        _font(p.add_run(f"N°  {numero}"), 10.5, color=C_MUTED)

        # Date — label au-dessus, valeur en dessous (même style CLIENT/AFFAIRE)
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, before=4, after=1)
        _font(p.add_run("DATE"), 7.5, bold=True, color=C_MUTED)

        p = rc.add_paragraph(d_cree)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, after=6)
        if p.runs:
            _font(p.runs[0], 9.5)

        # Client
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, before=8, after=2)
        _font(p.add_run("CLIENT"), 7.5, bold=True, color=C_MUTED)

        client = devis.client
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, after=1)
        _font(p.add_run(client.nom or "—"), 10.5, bold=True, color=C_NAVY)

        for line in [client.adresse,
                     f"{client.code_postal} {client.ville}".strip()]:
            if line.strip():
                p = rc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _spacing(p, after=1)
                if p.runs:
                    _font(p.runs[0], 9.5)

        affaire = (devis.reference_affaire or "").strip()
        if affaire:
            p = rc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _spacing(p, before=DOCX_HEADER_GAP_AFTER_CLIENT_PT, after=2)
            _font(p.add_run("AFFAIRE"), 7.5, bold=True, color=C_MUTED)
            p = rc.add_paragraph(affaire)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _spacing(p, after=DOCX_HEADER_GAP_AFTER_AFFAIRE_PT)
            if p.runs:
                _font(p.runs[0], 9.5)

        # Filet or sous l'en-tête
        sep_p = doc.add_paragraph()
        _spacing(sep_p, before=4, after=4)
        _bottom_border_para(sep_p, DOCX_FOOTER_LINE_COLOR, size="8")

        # ── Tableau principal ─────────────────────────────────────────────
        p_gap = doc.add_paragraph("")
        _spacing(p_gap, after=DOCX_SPACE_BEFORE_TABLE_PT)

        cols_count = len(LINE_TABLE_HEADERS)  # 5
        table = doc.add_table(rows=1, cols=cols_count)
        table.autofit = False
        for i, w in enumerate(DOCX_LINE_TABLE_WIDTHS_CM):
            table.columns[i].width = Cm(w)
        _fix_table(table, DOCX_CONTENT_TABLE_WIDTH_CM)

        # En-tête du tableau
        for i, title in enumerate(LINE_TABLE_HEADERS):
            cell = table.rows[0].cells[i]
            cell.text = title
            _shading(cell, HEX_TH_BG)
            _borders(cell, color=HEX_GRID, size="4")
            _margins(cell, 80, 80, 60, 60)
            p = cell.paragraphs[0]
            # col 0 = LEFT, cols 1+ = RIGHT (Qté, Unité, PU HT, Total HT)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 1 else WD_ALIGN_PARAGRAPH.LEFT
            _spacing(p, before=2, after=2)
            if p.runs:
                _font(p.runs[0], 9.5, bold=True, color=C_TH_TEXT)

        has_lots = self._has_grouped_lots(devis)

        if has_lots:
            for lot_idx, lot in enumerate(devis.lots, start=1):
                if not lot.lignes:
                    continue
                lot_name = lot.nom or f"Lot {lot_idx}"

                # En-tête lot : fond discret, merge all cols, texte ardoise gras
                lot_row = table.add_row()
                _no_split(lot_row)
                merged = lot_row.cells[0].merge(lot_row.cells[-1])
                merged.text = lot_name
                _shading(merged, HEX_LOT_BG)
                _borders(merged, color=HEX_GRID, size="4")
                _margins(merged, 80, 80, 60, 60)
                if merged.paragraphs[0].runs:
                    _font(merged.paragraphs[0].runs[0], 10, bold=True, color=C_NAVY)
                _spacing(merged.paragraphs[0], before=3, after=3)

                for line_idx, ligne in enumerate(lot.lignes):
                    row = table.add_row().cells
                    bg = HEX_ALT_BG if line_idx % 2 == 1 else "FFFFFF"
                    for c in row:
                        _shading(c, bg)
                        _borders(c, color=HEX_GRID, size="4",
                                 sides=("bottom",))
                        _margins(c, 80, 80, 55, 55)
                    _set_cell_multiline(row[0], ligne.designation or "")
                    row[1].text = "1" if ligne.unite.lower() == "forfait" else str(ligne.quantite)
                    row[2].text = "" if ligne.unite.lower() == "forfait" else ligne.unite
                    row[3].text = euro_fr(ligne.prix_unitaire_ht)
                    row[4].text = euro_fr(ligne.calculer_total_ht())
                    for ci in range(cols_count):
                        p = row[ci].paragraphs[0]
                        p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                                       if ci >= 1 else WD_ALIGN_PARAGRAPH.LEFT)
                        _spacing(p, after=2)
                        if p.runs:
                            _font(p.runs[0], 9.5)

                # Sous-total lot
                sub_row = table.add_row().cells
                for c in sub_row:
                    _shading(c, HEX_SUB_BG)
                    _borders(c, color=HEX_GRID, size="4")
                    _margins(c, 80, 80, 55, 55)
                sub_row[0].text = f"Sous-total  {lot_name}"
                sub_row[4].text = euro_fr(lot.calculer_sous_total_ht())
                sub_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for c in (sub_row[0], sub_row[4]):
                    _spacing(c.paragraphs[0], before=2, after=2)
                    if c.paragraphs[0].runs:
                        _font(c.paragraphs[0].runs[0], 9.5, bold=True, color=C_SUB_TXT)
        else:
            lines = self._iter_all_lines(devis)
            if not lines:
                row = table.add_row().cells
                for c in row:
                    _shading(c, "FFFFFF")
                    _borders(c, color=HEX_GRID, size="4", sides=("bottom",))
                row[0].text = "Aucune prestation"
                row[4].text = euro_fr(0)
                row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                for line_idx, ligne in enumerate(lines):
                    row = table.add_row().cells
                    bg = HEX_ALT_BG if line_idx % 2 == 1 else "FFFFFF"
                    for c in row:
                        _shading(c, bg)
                        _borders(c, color=HEX_GRID, size="4", sides=("bottom",))
                        _margins(c, 80, 80, 55, 55)
                    _set_cell_multiline(row[0], ligne.designation or "")
                    row[1].text = "1" if ligne.unite.lower() == "forfait" else str(ligne.quantite)
                    row[2].text = "" if ligne.unite.lower() == "forfait" else ligne.unite
                    row[3].text = euro_fr(ligne.prix_unitaire_ht)
                    row[4].text = euro_fr(ligne.calculer_total_ht())
                    for ci in range(cols_count):
                        p = row[ci].paragraphs[0]
                        p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                                       if ci >= 1 else WD_ALIGN_PARAGRAPH.LEFT)
                        _spacing(p, after=2)
                        if p.runs:
                            _font(p.runs[0], 9.5)

        # ── Banque (gauche) + Totaux (droite) sur la même ligne ─────────────
        # 4 colonnes : [banque | spacer | label totaux | valeur totaux]
        # cellule banque fusionnée verticalement sur les 3 lignes
        doc.add_paragraph("")

        totals_lbl_w = 4.6
        totals_val_w = 3.4   # total bloc = 8.0 cm, calé à droite
        BANK_W       = 7.0   # bloc banque compact à gauche
        SPACER_W     = DOCX_CONTENT_TABLE_WIDTH_CM - BANK_W - totals_lbl_w - totals_val_w

        totals_data = [
            ("Total HT",                            euro_fr(devis.calculer_total_ht()),  False),
            (f"TVA  {devis.tva_pourcent_global} %", euro_fr(devis.calculer_total_tva()), False),
            ("TOTAL TTC",                           euro_fr(devis.calculer_total_ttc()), True),
        ]

        def _vmerge_start(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            vm = OxmlElement("w:vMerge")
            vm.set(qn("w:val"), "restart")
            tc_pr.append(vm)

        def _vmerge_cont(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            vm = OxmlElement("w:vMerge")
            tc_pr.append(vm)

        bt = doc.add_table(rows=3, cols=4)
        bt.autofit = False
        bt.columns[0].width = Cm(BANK_W)
        bt.columns[1].width = Cm(SPACER_W)
        bt.columns[2].width = Cm(totals_lbl_w)
        bt.columns[3].width = Cm(totals_val_w)
        _fix_table(bt, DOCX_CONTENT_TABLE_WIDTH_CM)

        for r_i, (lbl_txt, val_txt, is_ttc) in enumerate(totals_data):
            cells = bt.rows[r_i].cells

            # ── Colonne 0 : banque (fusionnée verticalement) — largeur forcée
            _fix_cell_w(cells[0], BANK_W)
            if r_i == 0:
                if is_facture:
                    _vmerge_start(cells[0])
                    _shading(cells[0], HEX_SECTION)
                    _borders(cells[0], color=HEX_GRID, size="4")
                    _margins(cells[0], 120, 120, 100, 100)
                    cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cells[0].paragraphs[0].clear()
                    bank_lines = [
                        ("Coordonnées bancaires", ""),
                        ("", company.banque_nom),
                        ("", f"Code Banque {company.code_banque}  •  Code Guichet {company.code_guichet}"),
                        ("", f"IBAN : {company.iban}"),
                        ("", f"BIC : {company.bic}"),
                    ]
                    for idx, (lbl, val) in enumerate(bank_lines):
                        p = cells[0].paragraphs[0] if idx == 0 else cells[0].add_paragraph()
                        _spacing(p, after=2)
                        if lbl:
                            _font(p.add_run(f"{lbl}  " if val else lbl), 8.5, bold=True, color=C_NAVY)
                        if val:
                            _font(p.add_run(val), 9, color=C_TEXT)
                else:
                    _no_borders(cells[0])
                    _shading(cells[0], "FFFFFF")
            else:
                if is_facture:
                    _vmerge_cont(cells[0])
                    _shading(cells[0], HEX_SECTION)
                    _borders(cells[0], color=HEX_GRID, size="4", sides=("left", "right", "bottom"))
                else:
                    _no_borders(cells[0])
                    _shading(cells[0], "FFFFFF")

            # ── Colonne 1 : spacer transparent — largeur forcée
            _fix_cell_w(cells[1], SPACER_W)
            _no_borders(cells[1])
            _shading(cells[1], "FFFFFF")
            _margins(cells[1], 0, 0, 0, 0)

            # ── Colonnes 2 + 3 : label et valeur totaux — largeurs forcées
            _fix_cell_w(cells[2], totals_lbl_w)
            _fix_cell_w(cells[3], totals_val_w)
            hex_bg = HEX_NAVY if is_ttc else "FFFFFF"
            for c in (cells[2], cells[3]):
                _shading(c, hex_bg)
                _borders(c, color=HEX_BORDER, size="4")
                _margins(c, 80, 80, 60, 60)
            cells[2].text = lbl_txt
            cells[3].text = val_txt
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for p in (cells[2].paragraphs[0], cells[3].paragraphs[0]):
                _spacing(p, before=2, after=2)
            for c in (cells[2], cells[3]):
                if c.paragraphs[0].runs:
                    _font(c.paragraphs[0].runs[0],
                          11.5 if is_ttc else 9.5, bold=is_ttc,
                          color=C_WHITE if is_ttc else C_MUTED)
            if cells[3].paragraphs[0].runs:
                _font(cells[3].paragraphs[0].runs[0],
                      11.5 if is_ttc else 9.5, bold=is_ttc,
                      color=C_WHITE if is_ttc else None)

        # ── Notes / délais (en dessous, pleine largeur) ───────────────────
        notes_lines: list[tuple] = []
        if devis.delais.strip():
            notes_lines.append(("Délais", devis.delais))
        if devis.remarques.strip():
            notes_lines.append(("Remarques", devis.remarques))

        if notes_lines:
            doc.add_paragraph("")
            ment_tbl = doc.add_table(rows=1, cols=1)
            ment_tbl.autofit = False
            ment_tbl.columns[0].width = Cm(DOCX_CONTENT_TABLE_WIDTH_CM)
            _fix_table(ment_tbl, DOCX_CONTENT_TABLE_WIDTH_CM)
            mc = ment_tbl.cell(0, 0)
            _shading(mc, HEX_SECTION)
            _borders(mc, color=HEX_GRID, size="4")
            _margins(mc, 120, 120, 100, 100)
            mc.text = ""
            for idx, (lbl, val) in enumerate(notes_lines):
                p = mc.paragraphs[0] if idx == 0 else mc.add_paragraph()
                _spacing(p, after=2)
                if lbl:
                    _font(p.add_run(f"{lbl}  " if val else lbl), 8.5, bold=True, color=C_NAVY)
                if val:
                    _font(p.add_run(val), 9, color=C_TEXT)

        # ── Footer ────────────────────────────────────────────────────────
        footer_line1 = (f"{company.forme} {company.raison_sociale}  •  "
                        f"{company.adresse}  •  {company.code_postal_ville}")
        footer_line2 = (f"Tél. {company.telephone}  •  "
                        f"{company.email}  •  SIRET {company.siret}")

        footer = section.footer
        p1 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p1.text = ""
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p1, before=4, after=DOCX_FOOTER_LINE1_AFTER_PT, line=1.0)
        _top_border_para(p1, FOOTER_LINE, size="6")
        _font(p1.add_run(footer_line1), DOCX_FOOTER_LINE1_FONT_PT, color=FOOTER_TEXT)

        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p2, after=0, line=1.0)
        _font(p2.add_run(footer_line2), DOCX_FOOTER_LINE2_FONT_PT, color=FOOTER_TEXT)

        doc.save(output_path)
